
import io
import re
import uuid
from copy import deepcopy
from flask import Flask, request, jsonify, send_file, render_template
from flask_cors import CORS
from docx import Document

app = Flask(__name__, static_folder='static', template_folder='templates')
CORS(app)

SESSIONS = {}

BRACKETED_NAMED = re.compile(r"\[\s*([A-Za-z0-9][A-Za-z0-9 _.\-]{0,120})\s*\]")
BRACKETED_UNDERSCORES = re.compile(r"\[\s*_{2,}\s*\]")
QUOTED_TERM = re.compile(r"[“\"]([^”\"]+)[”\"]")

def iter_paragraphs(doc):
    for i, p in enumerate(doc.paragraphs):
        yield ('p', (i,), p)
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    yield ('t', (ti, ri, ci, pi), p)

def get_full_text(doc):
    parts = []
    for kind, idx, p in iter_paragraphs(doc):
        parts.append(p.text or "")
    return "\n".join(parts)

def detect_named_placeholders(doc_text):
    names = set()
    for m in BRACKETED_NAMED.finditer(doc_text):
        key = " ".join(m.group(1).split())
        if key.strip("_") and not set(key) <= set("_"):
            names.add(key)
    return sorted(names, key=lambda s: s.lower())

def label_from_context(before, after):
    m = re.search(r"\(\s*the\s+[“\"]([^”\"]+)", after, flags=re.IGNORECASE)
    if m:
        return m.group(1).strip()
    m = QUOTED_TERM.search(before[-120:])
    if m:
        return m.group(1).strip()
    tokens = re.findall(r"[A-Za-z][A-Za-z\-]+", before[-60:])
    if tokens:
        guess = " ".join(tokens[-3:]).strip()
        if len(guess) >= 3:
            return guess
    return None

def detect_underscore_placeholders_with_context(doc_text):
    res = []
    for m in BRACKETED_UNDERSCORES.finditer(doc_text):
        start, end = m.span()
        before = doc_text[max(0, start-160):start]
        after = doc_text[end:end+160]
        label = label_from_context(before, after)
        if not label:
            label = f"Blank {len(res)+1}"
        base = label
        k = 2
        while label in res:
            label = f"{base} ({k})"
            k += 1
        res.append(label)
    return res

def detect_signature_keys(doc):
    party = None
    keys = []
    for _, _, p in iter_paragraphs(doc):
        t = (p.text or "").strip()
        t_upper = t.upper().strip(":")
        if t_upper == "[COMPANY]" or t_upper == "COMPANY":
            party = "Company"; continue
        if t_upper.startswith("INVESTOR"):
            party = "Investor"; continue
        m = re.match(r"^(By|Name|Title|Address|Email):\s*$", t, flags=re.IGNORECASE)
        if m and party:
            label = m.group(1).capitalize()
            key = f"{party} {label}"
            if key not in keys:
                keys.append(key)
    return keys

def merge_unique(*lists):
    seen, out = set(), []
    for lst in lists:
        for item in lst:
            if item not in seen:
                seen.add(item); out.append(item)
    return out

def build_sessions_for_doc(original_bytes):
    doc = Document(io.BytesIO(original_bytes))
    txt = get_full_text(doc)
    named = detect_named_placeholders(txt)
    underscore = detect_underscore_placeholders_with_context(txt)
    sig_keys = detect_signature_keys(doc)
    all_keys = merge_unique(named, underscore, sig_keys)
    return {
        "original_doc_bytes": original_bytes,
        "named_keys": named,
        "underscore_keys": underscore,
        "signature_keys": sig_keys,
        "all_keys": all_keys,
        "answers": {}
    }

def sequential_replace_underscores(text, underscore_values):
    def repl(match):
        return underscore_values.pop(0) if underscore_values else ""
    return BRACKETED_UNDERSCORES.sub(lambda m: repl(m), text)

def replace_doc_content(doc, answers, session):
    named_replacements = {}
    for key in session["named_keys"]:
        if key in answers:
            named_replacements[f"[{key}]"] = answers[key]
            named_replacements[f"[ {key} ]"] = answers[key]

    under_values = [answers.get(k, "") for k in session["underscore_keys"]]

    for kind, idx, p in iter_paragraphs(doc):
        original = p.text or ""
        new_text = original
        for needle, val in named_replacements.items():
            if needle in new_text:
                new_text = new_text.replace(needle, val)
        if BRACKETED_UNDERSCORES.search(new_text):
            new_text = sequential_replace_underscores(new_text, under_values)
        if new_text != original:
            if p.runs:
                for _ in range(len(p.runs) - 1):
                    p._element.remove(p.runs[0]._element)
                p.runs[0].text = new_text
            else:
                p.add_run(new_text)

    # Signature fill pass
    flat = list(iter_paragraphs(doc))
    for i, (kind, idx, p) in enumerate(flat):
        t = (p.text or "").strip()
        m = re.match(r"^(By|Name|Title|Address|Email):\s*$", t, flags=re.IGNORECASE)
        if not m: continue
        party = None
        for j in range(i-1, max(-1, i-12), -1):
            tj = (flat[j][2].text or "").strip().upper().strip(":")
            if tj in ("[COMPANY]", "COMPANY"):
                party = "Company"; break
            if tj.startswith("INVESTOR"):
                party = "Investor"; break
        if not party: continue
        label = m.group(1).capitalize()
        key = f"{party} {label}"
        value = answers.get(key)
        if value is None: continue
        lines = str(value).splitlines() or [""]
        # next paragraph
        next_p = flat[i+1][2] if i+1 < len(flat) else None
        if next_p is None:
            continue
        # replace next paragraph with first line
        if next_p.runs:
            for _ in range(len(next_p.runs) - 1):
                next_p._element.remove(next_p.runs[0]._element)
            next_p.runs[0].text = lines[0]
        else:
            next_p.add_run(lines[0])
        # for extra lines, append new paragraphs after next_p
        from docx.oxml.text.paragraph import CT_P
        from docx.text.paragraph import Paragraph
        anchor = next_p._p
        parent = next_p._parent
        for extra in lines[1:]:
            new_ct_p = CT_P()
            anchor.addnext(new_ct_p)
            new_par = Paragraph(new_ct_p, parent)
            new_par.add_run(extra)
            anchor = new_ct_p

    return doc

def make_preview_html(doc):
    html_lines = []
    for kind, idx, p in iter_paragraphs(doc):
        txt = (p.text or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        html_lines.append(f"<p>{txt or '&nbsp;'}</p>")
    return "<div>" + "\n".join(html_lines) + "</div>"

@app.route("/", methods=["GET"])
def home():
    return render_template("index.html")

@app.route("/api/health", methods=["GET"])
def health():
    return jsonify({"ok": True})

@app.route("/api/upload", methods=["POST"])
def upload():
    if "file" not in request.files:
        return jsonify({"error": "Missing file"}), 400
    f = request.files["file"]
    if not f.filename.lower().endswith(".docx"):
        return jsonify({"error": "Only .docx is supported"}), 400
    original_bytes = f.read()
    try:
        sess = build_sessions_for_doc(original_bytes)
    except Exception as e:
        return jsonify({"error": f"Failed to read .docx: {e}"}), 400
    session_id = str(uuid.uuid4())
    SESSIONS[session_id] = sess
    return jsonify({
        "session_id": session_id,
        "placeholders": sess["all_keys"],
        "meta": {
            "named_keys": sess["named_keys"],
            "underscore_keys": sess["underscore_keys"],
            "signature_keys": sess["signature_keys"]
        }
    })

@app.route("/api/placeholders", methods=["GET"])
def get_placeholders():
    session_id = request.args.get("session_id", "")
    sess = SESSIONS.get(session_id)
    if not sess:
        return jsonify({"error": "Invalid session"}), 404
    return jsonify({
        "placeholders": sess["all_keys"],
        "answered": list(sess["answers"].keys()),
        "meta": {
            "named_keys": sess["named_keys"],
            "underscore_keys": sess["underscore_keys"],
            "signature_keys": sess["signature_keys"]
        }
    })

@app.route("/api/answer", methods=["POST"])
def post_answer():
    data = request.get_json(silent=True) or {}
    session_id = data.get("session_id")
    field = data.get("field")
    value = data.get("value", "")
    sess = SESSIONS.get(session_id)
    if not sess:
        return jsonify({"error": "Invalid session"}), 404
    if field not in sess["all_keys"]:
        return jsonify({"error": f"Unknown field '{field}'"}), 400
    sess["answers"][field] = str(value)
    remaining = [p for p in sess["all_keys"] if p not in sess["answers"]]
    return jsonify({"saved": True, "next": (remaining[0] if remaining else None), "remaining_count": len(remaining)})

@app.route("/api/preview", methods=["GET"])
def preview():
    session_id = request.args.get("session_id", "")
    sess = SESSIONS.get(session_id)
    if not sess:
        return jsonify({"error": "Invalid session"}), 404
    doc = Document(io.BytesIO(sess["original_doc_bytes"]))
    doc = replace_doc_content(doc, sess["answers"], sess)
    html = make_preview_html(doc)
    return html, 200, {"Content-Type": "text/html; charset=utf-8"}

@app.route("/api/download", methods=["GET"])
def download():
    session_id = request.args.get("session_id", "")
    sess = SESSIONS.get(session_id)
    if not sess:
        return jsonify({"error": "Invalid session"}), 404
    doc = Document(io.BytesIO(sess["original_doc_bytes"]))
    doc = replace_doc_content(doc, sess["answers"], sess)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return send_file(out, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name="completed_document.docx")

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000, debug=True)
